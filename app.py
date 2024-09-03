import streamlit as st
import pyrebase
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import openai
import json
import uuid
import PyPDF2
import docx
import io
import logging
import networkx as nx
import base64
from web3 import Web3
import os
from dotenv import load_dotenv
import requests
import time
import pandas as pd
from datetime import datetime
import folium
from streamlit_folium import folium_static
from geopy.geocoders import Nominatim
from geopy.exc import GeocoderTimedOut, GeocoderUnavailable
import pandas as pd
import plotly.express as px
import difflib
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import seaborn as sns
from folium.plugins import HeatMap
from io import BytesIO
import reportlab
from exa_py import Exa
from datetime import datetime, timedelta


load_dotenv()

# Pinata setup
PINATA_API_KEY = os.getenv("PINATA_API_KEY")
PINATA_SECRET_API_KEY = os.getenv("PINATA_SECRET_API_KEY")
PINATA_JWT = os.getenv("PINATA_JWT")
PINATA_BASE_URL = "https://api.pinata.cloud"
EXA_API_KEY = os.getenv("EXA_API_KEY")

# Web3 setup
WEB3_PROVIDER_URL = os.getenv("WEB3_PROVIDER_URL")
web3 = Web3(Web3.HTTPProvider(WEB3_PROVIDER_URL))

# Set up OpenAI client
AI71_BASE_URL = "https://api.ai71.ai/v1/"
AI71_API_KEY = ""
client = openai.OpenAI(
    api_key=AI71_API_KEY,
    base_url=AI71_BASE_URL,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Firebase Configuration
firebaseConfig = {
    "apiKey": "AIzaSyC-G1JRKBZtWjtu02WKCmizdBmXMMP3PAA",
    "authDomain": "investiboard.firebaseapp.com",
    "projectId": "investiboard",
    "storageBucket": "investiboard.appspot.com",
    "messagingSenderId": "556809275550",
    "appId": "1:556809275550:web:a37e3770466ea4011b138f",
    "measurementId": "G-DGNP1CF7QK",
    "databaseURL" : "https://investiboard-default-rtdb.asia-southeast1.firebasedatabase.app/"
}

# Initialize Firebase
try:
    firebase = pyrebase.initialize_app(firebaseConfig)
    auth = firebase.auth()
    db = firebase.database()
except Exception as e:
    st.error(f"Error initializing Firebase: {str(e)}")
    st.stop()

vectorizer = TfidfVectorizer()

def initialize_session_state():
    if 'user' not in st.session_state:
        st.session_state.user = None
    if "response_added" not in st.session_state:
        st.session_state.response_added = False
    if 'authentication_status' not in st.session_state:
        st.session_state.authentication_status = None
    if 'cards' not in st.session_state:
        st.session_state.cards = []
    if 'connections' not in st.session_state:
        st.session_state.connections = []
    if 'graph' not in st.session_state:
        st.session_state.graph = nx.Graph()
    if 'card_updates' not in st.session_state:
        st.session_state.card_updates = None
    if 'show_rag' not in st.session_state:
        st.session_state.show_rag = False
    if 'card_just_created' not in st.session_state:
        st.session_state.card_just_created = False
    if 'wallet_connected' not in st.session_state:
        st.session_state.wallet_connected = False
    if 'wallet_address' not in st.session_state:
        st.session_state.wallet_address = None
    if 'show_map' not in st.session_state:
        st.session_state.show_map = False


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def connect_wallet():
    st.session_state.wallet_address = st.text_input("Enter your Ethereum wallet address:")
    if st.session_state.wallet_address and web3.isAddress(st.session_state.wallet_address):
        st.session_state.wallet_connected = True
        st.success(f"Wallet connected: {st.session_state.wallet_address}")
    else:
        st.error("Invalid Ethereum address. Please enter a valid address.")
        st.session_state.wallet_connected = False
        st.session_state.wallet_address = None

def disconnect_wallet():
    st.session_state.wallet_connected = False
    st.session_state.wallet_address = None
    st.success("Wallet disconnected successfully.")

def upload_to_ipfs(file):
    try:
        files = {
            'file': file.getvalue()
        }
        headers = {
            'Authorization': f'Bearer {PINATA_JWT}'
        }
        response = requests.post(
            PINATA_BASE_URL + "/pinning/pinFileToIPFS",
            files=files,
            headers=headers
        )
        if response.status_code == 200:
            return response.json()['IpfsHash']
        else:
            st.error(f"Error uploading to IPFS: {response.text}")
            return None
    except Exception as e:
        st.error(f"Error uploading to IPFS: {str(e)}")
        return None

def get_ipfs_files():
    try:
        headers = {
            'Authorization': f'Bearer {PINATA_JWT}'
        }
        response = requests.get(
            PINATA_BASE_URL + "/data/pinList?status=pinned",
            headers=headers
        )
        if response.status_code == 200:
            return response.json()['rows']
        else:
            st.error(f"Error fetching IPFS files: {response.text}")
            return []
    except Exception as e:
        st.error(f"Error fetching IPFS files: {str(e)}")
        return []
    
def sign_up():
    st.subheader("Create New Account")
    with st.form("signup_form"):
        new_user_email = st.text_input("Email Address", key="signup_email")
        new_user_password = st.text_input("Password", type='password', key="signup_password")
        submit_button = st.form_submit_button("Sign Up")
    
    if submit_button:
        try:
            user = auth.create_user_with_email_and_password(new_user_email, new_user_password)
            st.success("Account created successfully!")
            st.session_state.user = user
            st.session_state.authentication_status = True
        except Exception as e:
            st.error(f"Error: {e}")

def sign_in():
    st.subheader("Sign In to Existing Account")
    with st.form("signin_form"):
        user_email = st.text_input("Email Address", key="signin_email")
        user_password = st.text_input("Password", type='password', key="signin_password")
        submit_button = st.form_submit_button("Sign In")
    
    if submit_button:
        try:
            user = auth.sign_in_with_email_and_password(user_email, user_password)
            st.success("Signed in successfully!")
            st.session_state.user = user
            st.session_state.authentication_status = True
            load_canvas_data()  
        except Exception as e:
            st.error(f"Error: {e}")

def sign_out():
    st.session_state.user = None
    st.session_state.authentication_status = None
    st.session_state.cards = []
    st.session_state.connections = []
    st.session_state.graph = nx.Graph()
    st.success("Signed out successfully!")


def google_sign_in():
    st.subheader("Sign In with Google")
    auth_url = f"https://accounts.google.com/o/oauth2/v2/auth?client_id={firebaseConfig['clientId']}&redirect_uri={firebaseConfig['redirectUri']}&response_type=code&scope=email%20profile"
    
    if st.button("Sign In with Google"):
        st.markdown(f'<a href="{auth_url}" target="_self">Click here to sign in with Google</a>', unsafe_allow_html=True)

def save_canvas_data(canvas_data):
    if st.session_state.user:
        user_id = st.session_state.user['localId']
        db.child("users").child(user_id).child("canvas").set(canvas_data)
        st.session_state.cards = canvas_data["cards"]
        st.session_state.connections = canvas_data["connections"]

def load_canvas_data():
    if st.session_state.user:
        user_id = st.session_state.user['localId']
        canvas_data = db.child("users").child(user_id).child("canvas").get().val()
        if canvas_data is None:
            canvas_data = {"cards": [], "connections": []}
        
        st.session_state.cards = canvas_data.get("cards", [])
        st.session_state.connections = canvas_data.get("connections", [])
        
        return canvas_data
    return {"cards": [], "connections": []}



def create_event_card():
    with st.form("event_card_form"):
        title = st.text_input("Title/Name")
        date = st.date_input("Date")
        time = st.time_input("Time")
        location = st.text_input("Location")
        description = st.text_area("Description")
        related_suspects = st.multiselect("Related Suspects", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Suspect'])
        related_evidence = st.multiselect("Related Evidence", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Evidence'])
        notes = st.text_area("Notes")
        submitted = st.form_submit_button("Add Event")
    
    if submitted:
        if not all([title, date, time, location, description]):
            st.error("Please fill in all required fields.")
            return None
        return {
            "id": str(uuid.uuid4()),
            "type": "Event",
            "title": title,
            "date_time": datetime.combine(date, time).isoformat(),
            "location": location,
            "description": description,
            "related_suspects": related_suspects,
            "related_evidence": related_evidence,
            "notes": notes,
            "x": 50,
            "y": len(st.session_state.cards) * 220
        }
    return None

def create_suspect_card():
    with st.form("suspect_card_form"):
        name = st.text_input("Name")
        photo = st.file_uploader("Photo", type=["jpg", "jpeg", "png"])
        dob = st.date_input("Date of Birth")
        occupation = st.text_input("Occupation")
        address = st.text_input("Address")
        alibi = st.text_area("Alibi")
        description = st.text_area("Description")
        related_events = st.multiselect("Related Events", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Event'])
        related_evidence = st.multiselect("Related Evidence", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Evidence'])
        notes = st.text_area("Notes")
        submitted = st.form_submit_button("Add Suspect")
    
    if submitted:
        if not all([name, dob, occupation, address, description]):
            st.error("Please fill in all required fields.")
            return None
        suspect_card = {
            "id": str(uuid.uuid4()),
            "type": "Suspect",
            "name": name,
            "dob": dob.isoformat(),
            "occupation": occupation,
            "address": address,
            "alibi": alibi,
            "description": description,
            "related_events": related_events,
            "related_evidence": related_evidence,
            "notes": notes,
            "x": 50,
            "y": len(st.session_state.cards) * 220
        }
        if photo:
            suspect_card["photo"] = base64.b64encode(photo.read()).decode()
        return suspect_card
    return None

def create_evidence_card():
    with st.form("evidence_card_form"):
        title = st.text_input("Title/Name")
        evidence_type = st.text_input("Type")
        date = st.date_input("Date")
        time = st.time_input("Time")
        location = st.text_input("Location")
        description = st.text_area("Description")
        related_events = st.multiselect("Related Events", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Event'])
        related_suspects = st.multiselect("Related Suspects", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Suspect'])
        condition = st.text_input("Condition")
        notes = st.text_area("Notes")
        submitted = st.form_submit_button("Add Evidence")
    
    if submitted:
        if not all([title, evidence_type, date, time, location, description, condition]):
            st.error("Please fill in all required fields.")
            return None
        return {
            "id": str(uuid.uuid4()),
            "type": "Evidence",
            "title": title,
            "evidence_type": evidence_type,
            "date_time": datetime.combine(date, time).isoformat(),
            "location": location,
            "description": description,
            "related_events": related_events,
            "related_suspects": related_suspects,
            "condition": condition,
            "notes": notes,
            "x": 50,
            "y": len(st.session_state.cards) * 220
        }
    return None

def create_card(card_type):
    with st.form(key=f"create_{card_type.lower()}_card"):
        st.subheader(f"Create {card_type} Card")
        new_card = {"id": str(uuid.uuid4()), "type": card_type, "x": 50, "y": 50}
        
        if card_type == "Event":
            new_card["title"] = st.text_input("Title/Name")
            new_card["date"] = st.date_input("Date").isoformat()
            new_card["time"] = st.time_input("Time").isoformat()
            new_card["location"] = st.text_input("Location")
            new_card["description"] = st.text_area("Description")
            new_card["related_suspects"] = st.multiselect("Related Suspects", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Suspect'])
            new_card["related_evidence"] = st.multiselect("Related Evidence", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Evidence'])
            new_card["notes"] = st.text_area("Notes")
            required_fields = ["title", "date", "time", "location", "description"]
        
        elif card_type == "Suspect":
            new_card["name"] = st.text_input("Name")
            new_card["dob"] = st.date_input("Date of Birth").isoformat()
            new_card["occupation"] = st.text_input("Occupation")
            new_card["address"] = st.text_input("Address")
            new_card["alibi"] = st.text_area("Alibi")
            new_card["description"] = st.text_area("Description")
            new_card["related_events"] = st.multiselect("Related Events", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Event'])
            new_card["related_evidence"] = st.multiselect("Related Evidence", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Evidence'])
            new_card["notes"] = st.text_area("Notes")
            required_fields = ["name", "dob", "occupation", "address", "description"]
        
        elif card_type == "Evidence":
            new_card["title"] = st.text_input("Title/Name")
            new_card["evidence_type"] = st.text_input("Evidence Type")
            new_card["date"] = st.date_input("Date Found").isoformat()
            new_card["time"] = st.time_input("Time Found").isoformat()
            new_card["location"] = st.text_input("Location")
            new_card["description"] = st.text_area("Description")
            new_card["related_events"] = st.multiselect("Related Events", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Event'])
            new_card["related_suspects"] = st.multiselect("Related Suspects", [get_card_identifier(card) for card in st.session_state.cards if card['type'] == 'Suspect'])
            new_card["condition"] = st.text_input("Condition")
            new_card["notes"] = st.text_area("Notes")
            required_fields = ["title", "evidence_type", "date", "location", "description"]


        uploaded_file = st.file_uploader("Upload an image", type=["jpg", "jpeg", "png"])
        if uploaded_file is not None:
            image_data = uploaded_file.read()
            new_card["image"] = base64.b64encode(image_data).decode()
        
        submitted = st.form_submit_button("Create Card")

        
        
        if submitted:
            if all(new_card.get(field) for field in required_fields):
                if "cards" not in st.session_state:
                    st.session_state.cards = []
                st.session_state.cards.append(new_card)
                save_canvas_data({"cards": st.session_state.cards, "connections": st.session_state.get("connections", [])})
                st.success(f"{card_type} card created successfully.")
                st.experimental_rerun()  
            else:
                st.error("Please fill in all required fields.")


def update_card_position(card_id, x, y):
    canvas_data = load_canvas_data()
    for card in canvas_data["cards"]:
        if card['id'] == card_id:
            card['x'] = x
            card['y'] = y
            break
    save_canvas_data(canvas_data)


def delete_card(card_id):
    canvas_data = load_canvas_data()
    canvas_data["cards"] = [card for card in canvas_data["cards"] if card['id'] != card_id]
    canvas_data["connections"] = [conn for conn in canvas_data["connections"] if conn[0] != card_id and conn[1] != card_id]
    save_canvas_data(canvas_data)
    st.success("Card deleted successfully.")


def edit_card():
    if 'edit_card_id' in st.session_state:
        card_id = st.session_state.edit_card_id
        card = next((card for card in st.session_state.cards if card['id'] == card_id), None)
        if card:
            with st.form(key=f"edit_card_{card_id}"):
                st.subheader(f"Edit {card['type']} Card")
                updated_card = {}
                for key, value in card.items():
                    if key not in ['id', 'type', 'x', 'y']:
                        updated_card[key] = st.text_input(key.capitalize(), value)
                
                if st.form_submit_button("Save Changes"):
                    card.update(updated_card)
                    update_card(card_id, card)
                    del st.session_state.edit_card_id
                    st.success("Card updated successfully.")
                    st.rerun()

def clear_all_cards():
    st.session_state.cards = []
    st.session_state.connections = []
    
    if st.session_state.user:
        user_id = st.session_state.user['localId']
        db.child("users").child(user_id).child("canvas").set({})
    
    st.success("All cards have been cleared from the board and database.")
    st.experimental_rerun()

def connect_cards():
    if len(st.session_state.cards) < 2:
        st.warning("You need at least two cards to make a connection.")
        return

    card_identifiers = [get_card_identifier(card) for card in st.session_state.cards]
    
    card1 = st.selectbox("Select first card", card_identifiers, key="card1")
    card2 = st.selectbox("Select second card", card_identifiers, key="card2")
    
    if st.button("Connect Cards"):
        if card1 != card2 and (card1, card2) not in st.session_state.connections and (card2, card1) not in st.session_state.connections:
            st.session_state.connections.append((card1, card2))
            save_canvas_data({"cards": st.session_state.cards, "connections": st.session_state.connections})
            st.success("Cards connected successfully!")
            st.experimental_rerun()
        else:
            st.error("Invalid connection or already exists.")   

def get_card_identifier(card):
    return card.get('title') or card.get('name') or 'Unnamed Card'

def get_card_identifier(card):
    return card.get('title') or card.get('name') or 'Unnamed Card'

def process_document(file, file_type):
    if file_type == 'pdf':
        reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
    elif file_type == 'docx':
        doc = docx.Document(io.BytesIO(file.read()))
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        text = file.getvalue().decode('utf-8')
    
    return text

def store_document(text, metadata):
    doc_id = str(uuid.uuid4())
    doc_data = {
        'text': text,
        'metadata': metadata
    }
    db.child('rag_documents').child(doc_id).set(doc_data)
    return doc_id

def retrieve_similar_documents(query_vector, top_k=3):
    all_docs = db.child('rag_documents').get().val()
    if not all_docs:
        return []
    
    all_texts = [doc_data['text'] for doc_data in all_docs.values()]
    all_vectors = vectorizer.fit_transform(all_texts)
    
    similarities = cosine_similarity(query_vector, all_vectors)[0]
    top_indices = similarities.argsort()[-top_k:][::-1]
    
    similar_docs = []
    for idx in top_indices:
        doc_id = list(all_docs.keys())[idx]
        similar_docs.append((doc_id, similarities[idx], all_docs[doc_id]))
    
    return similar_docs

def generate_response(query, similar_docs):
    context = "Here are the most relevant documents for the query:\n\n"
    for i, (doc_id, similarity, doc_data) in enumerate(similar_docs, 1):
        context += f"Document {i} (Similarity: {similarity:.2f}):\n"
        context += f"Title: {doc_data['metadata'].get('filename', 'Untitled')}\n"
        context += f"Content: {doc_data['text'][:500]}...\n\n"

    prompt = f"""Based on the following context, please answer the user's question. If the answer is not directly stated in the context, use the information provided to infer a response. If you cannot answer based on the given context, please say so.

Context:
{context}

User Question: {query}

Answer:"""

    try:
        response = client.chat.completions.create(
            model="tiiuae/falcon-180b-chat",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that answers questions based on the given context. Always refer to the provided documents when answering."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        logging.error(f"Error in generate_response: {str(e)}")
        return "I'm sorry, but I encountered an error while processing your request. Please try again later."

def rag_query(query):
    all_docs = db.child('rag_documents').get().val()
    if not all_docs:
        return "No documents have been uploaded yet. Please upload some documents before asking questions."
    
    all_texts = [doc_data['text'] for doc_data in all_docs.values()]
    vectorizer.fit(all_texts)
    
    query_vector = vectorizer.transform([query])
    
    similar_docs = retrieve_similar_documents(query_vector)
    
    response = generate_response(query, similar_docs)
    
    return response

def rag_interface():
    st.subheader("Document-based Question Answering System")
    tool_selection = st.radio("Select Tool", ["RAG", "IPDR Analysis"])

    if tool_selection == "RAG":

        uploaded_file = st.file_uploader("Upload a document (PDF, DOCX, or TXT)", type=['pdf', 'docx', 'txt'])
        if uploaded_file:
            try:
                file_type = uploaded_file.type.split('/')[-1]
                text = process_document(uploaded_file, file_type)
                metadata = {
                    'filename': uploaded_file.name,
                    'filetype': file_type
                }
                doc_id = store_document(text, metadata)
                st.success(f"Document uploaded and processed. ID: {doc_id}")
            except Exception as e:
                st.error(f"Error processing document: {str(e)}")
        
        query = st.text_input("Ask a question based on the uploaded documents:")
        if st.button("Get Answer"):
            if query:
                with st.spinner("Generating answer..."):
                    response = rag_query(query)
                st.write("Answer:", response)
            else:
                st.warning("Please enter a question.")

    elif tool_selection == "IPDR Analysis":
        ipdr_analysis()

def create_tabular_view(cards):
    df = pd.DataFrame(cards)
    
    if 'date' in df.columns:
        df['date'] = pd.to_datetime(df['date'], errors='coerce').dt.date
    else:
        df['date'] = pd.NaT

    if 'time' in df.columns:
        df['time'] = pd.to_datetime(df['time'], format='%H:%M:%S', errors='coerce').dt.time
    else:
        df['time'] = pd.NaT

    for col in ['location', 'name', 'evidence_type']:
        if col not in df.columns:
            df[col] = ''

    return df

def display_tabular_view(df, sort_by='date', ascending=True, filter_option=None, filter_value=None):
    if filter_option and filter_value:
        if filter_option in ['location', 'name', 'evidence_type']:
            df = df[df[filter_option].str.contains(filter_value, case=False, na=False)]
        elif filter_option == 'suspect_name':
            df = df[df['name'].str.contains(filter_value, case=False, na=False) & (df['type'] == 'Suspect')]
        elif filter_option == 'evidence_related':
            df = df[df['evidence_type'].str.contains(filter_value, case=False, na=False) | 
                    df['related_evidence'].astype(str).str.contains(filter_value, case=False, na=False)]

    if sort_by in df.columns:
        df_sorted = df.sort_values(by=sort_by, ascending=ascending)
    else:
        st.warning(f"Cannot sort by '{sort_by}'. Column not found in the data.")
        df_sorted = df

    return df_sorted

def display_draggable_cards():
    canvas_data = load_canvas_data()
    card_data = json.dumps(canvas_data.get("cards", []))
    connection_data = json.dumps(canvas_data.get("connections", []))
    
    html_content = f"""
    <style>
    #outer-container {{
        position: relative;
        width: 100%;
        height: 800px;
        overflow: hidden;
        border: 1px solid #ccc;
    }}
    #canvas-container {{
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background-color: #f0f0f0;
        cursor: move;
    }}
    #connectionCanvas {{
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        pointer-events: none;
        z-index: 1;
    }}
    #cardContainer {{
        position: absolute;
        top: 0;
        left: 0;
    }}

    .card-image {{
    max-width: 100%;
    height: auto;
    margin-top: 10px;
    }}

    .card {{
        position: absolute;
        width: 300px;
        min-height: 200px;
        padding: 15px;
        border: 1px solid #999;
        border-radius: 8px;
        font-family: Arial, sans-serif;
        font-size: 14px;
        cursor: move;
        opacity: 0.9;
        transition: box-shadow 0.3s ease;
        background-color: #ffffff;
        z-index: 2;
    }}
    .card:hover {{
        box-shadow: 0 0 15px rgba(0,0,0,0.2);
    }}
    .card-content {{
        display: table;
        width: 100%;
    }}
    .card-row {{
        display: table-row;
    }}
    .card-cell {{
        display: table-cell;
        padding: 5px;
        border-bottom: 1px solid #eee;
    }}
    .card-cell:first-child {{
        font-weight: bold;
        width: 30%;
    }}
    .delete-btn {{
        position: absolute;
        top: 10px;
        right: 10px;
        background-color: #ff4d4d;
        color: white;
        border: none;
        border-radius: 4px;
        padding: 5px 10px;
        cursor: pointer;
        transition: background-color 0.3s ease;
    }}
    .delete-btn:hover {{
        background-color: #ff1a1a;
    }}
    .card-image {{
        max-width: 100%;
        height: auto;
        margin-top: 10px;
    }}
    #mini-map {{
        position: absolute;
        bottom: 10px;
        right: 10px;
        width: 200px;
        height: 150px;
        background-color: rgba(255, 255, 255, 0.8);
        border: 1px solid #999;
        z-index: 3;
    }}
    #mini-map-canvas {{
        width: 100%;
        height: 100%;
    }}
    #zoom-controls {{
        position: absolute;
        top: 10px;
        left: 10px;
        z-index: 3;
    }}
    .zoom-btn {{
        width: 30px;
        height: 30px;
        font-size: 20px;
        margin: 5px;
    }}
    </style>
    <div id="outer-container">
        <div id="canvas-container">
            <canvas id="connectionCanvas"></canvas>
            <div id="cardContainer"></div>
        </div>
        <div id="mini-map">
            <canvas id="mini-map-canvas"></canvas>
        </div>
        <div id="zoom-controls">
            <button class="zoom-btn" onclick="zoomIn()">+</button>
            <button class="zoom-btn" onclick="zoomOut()">-</button>
        </div>
    </div>

    <script>
    const cardData = {card_data};
    const connectionData = {connection_data};
    const outerContainer = document.getElementById('outer-container');
    const canvasContainer = document.getElementById('canvas-container');
    const cardContainer = document.getElementById('cardContainer');
    const canvas = document.getElementById('connectionCanvas');
    const ctx = canvas.getContext('2d');
    const miniMap = document.getElementById('mini-map');
    const miniMapCanvas = document.getElementById('mini-map-canvas');
    const miniMapCtx = miniMapCanvas.getContext('2d');

    let scale = 1;
    let containerWidth = 5000;  // Increased initial width
    let containerHeight = 3000;  // Increased initial height
    let offsetX = 0;
    let offsetY = 0;

    function resizeCanvas() {{
        canvas.width = containerWidth;
        canvas.height = containerHeight;
        cardContainer.style.width = containerWidth + 'px';
        cardContainer.style.height = containerHeight + 'px';
        drawConnections();
        updateMiniMap();
    }}

    function zoomIn() {{
        scale = Math.min(scale * 1.2, 3);
        applyZoom();
    }}

    function zoomOut() {{
        scale = Math.max(scale / 1.2, 0.5);
        applyZoom();
    }}

    function applyZoom() {{
        cardContainer.style.transform = `scale(${{scale}})`;
        cardContainer.style.transformOrigin = '0 0';
        drawConnections();
        updateMiniMap();
    }}

    function updateMiniMap() {{
        miniMapCanvas.width = miniMap.offsetWidth;
        miniMapCanvas.height = miniMap.offsetHeight;
        miniMapCtx.clearRect(0, 0, miniMapCanvas.width, miniMapCanvas.height);

        const scaleX = miniMapCanvas.width / containerWidth;
        const scaleY = miniMapCanvas.height / containerHeight;

        // Draw cards
        cardData.forEach(card => {{
            const x = card.x * scaleX;
            const y = card.y * scaleY;
            miniMapCtx.fillStyle = 'blue';
            miniMapCtx.fillRect(x, y, 5, 5);
        }});

        // Draw visible area
        const visibleWidth = outerContainer.offsetWidth / scale;
        const visibleHeight = outerContainer.offsetHeight / scale;
        const visibleX = -offsetX / scale * scaleX;
        const visibleY = -offsetY / scale * scaleY;
        miniMapCtx.strokeStyle = 'red';
        miniMapCtx.strokeRect(visibleX, visibleY, visibleWidth * scaleX, visibleHeight * scaleY);
    }}

    function getCardIdentifier(card) {{
        return card.title || card.name || 'Unnamed Card';
    }}

    function initializeCards(cardData) {{
    cardContainer.innerHTML = '';  // Clear existing cards
    cardData.forEach(createCard);
    }}

    function createCard(card) {{
        if (document.getElementById(card.id)) {{
        return;
        }}
        const cardElement = document.createElement('div');
        cardElement.className = 'card';
        cardElement.id = card.id;
        cardElement.style.backgroundColor = card.type === 'Event' ? 'rgba(255,179,186,0.9)' : 
                                            card.type === 'Suspect' ? 'rgba(186,255,201,0.9)' : 'rgba(186,225,255,0.9)';
        cardElement.style.left = (card.x || 50) + 'px';
        cardElement.style.top = (card.y || 50) + 'px';
        
        let cardContent = `<div class="card-content">`;
        cardContent += `<div class="card-row"><div class="card-cell">Type:</div><div class="card-cell">${{card.type}}</div></div>`;
        for (const [key, value] of Object.entries(card)) {{
            if (!['id', 'type', 'x', 'y', 'image'].includes(key)) {{
                cardContent += `<div class="card-row"><div class="card-cell">${{key}}:</div><div class="card-cell">${{value}}</div></div>`;
            }}
        }}
        cardContent += `</div>`;
        
        if (card.image) {{
            cardContent += `<img src="data:image/jpeg;base64,${{card.image}}" class="card-image" alt="Card Image">`;
        }}
        
        cardElement.innerHTML = cardContent;
        cardElement.addEventListener('mousedown', startDragging);
        cardContainer.appendChild(cardElement);
    }}
    let isDragging = false;
    let isPanning = false;
    let currentCard = null;
    let startX, startY;

    function startDragging(e) {{
        isDragging = true;
        currentCard = e.currentTarget;
        const rect = currentCard.getBoundingClientRect();
        startX = e.clientX - rect.left;
        startY = e.clientY - rect.top;
        
        document.addEventListener('mousemove', onMouseMove);
        document.addEventListener('mouseup', stopDragging);
    }}

    function onMouseMove(e) {{
        if (isDragging) {{
            const containerRect = cardContainer.getBoundingClientRect();
            let newX = (e.clientX - containerRect.left - startX) / scale;
            let newY = (e.clientY - containerRect.top - startY) / scale;
            
            newX = Math.max(0, Math.min(newX, containerWidth - currentCard.offsetWidth));
            newY = Math.max(0, Math.min(newY, containerHeight - currentCard.offsetHeight));
            
            currentCard.style.left = newX + 'px';
            currentCard.style.top = newY + 'px';
            
            drawConnections();
        }}
    }}

    function stopDragging() {{
        if (isDragging) {{
            isDragging = false;
            updateCardPosition(currentCard.id, parseFloat(currentCard.style.left), parseFloat(currentCard.style.top));
            drawConnections();
        }}
        currentCard = null;
        document.removeEventListener('mousemove', onMouseMove);
        document.removeEventListener('mouseup', stopDragging);
    }}

    function updateCardPosition(cardId, x, y) {{
        const cardIndex = cardData.findIndex(c => c.id === cardId);
        if (cardIndex !== -1) {{
            cardData[cardIndex].x = x;
            cardData[cardIndex].y = y;
            updateStreamlit({{ type: 'update', cardId: cardId, x: x, y: y }});
            drawConnections();
        }}
    }}


    function drawConnections() {{
        ctx.clearRect(0, 0, canvas.width, canvas.height);
        ctx.strokeStyle = 'rgba(0, 123, 255, 0.7)';
        ctx.lineWidth = 10;

        connectionData.forEach(([c1, c2]) => {{
            const card1 = cardData.find(c => getCardIdentifier(c) === c1);
            const card2 = cardData.find(c => getCardIdentifier(c) === c2);
            if (card1 && card2) {{
                const elem1 = document.getElementById(card1.id);
                const elem2 = document.getElementById(card2.id);
                if (elem1 && elem2) {{
                    const pos1 = getCardPosition(elem1);
                    const pos2 = getCardPosition(elem2);

                    drawArrow(pos1.x, pos1.y, pos2.x, pos2.y);
                }}
            }}
        }});
    }}

    function getCardPosition(element) {{
        const rect = element.getBoundingClientRect();
        const canvasRect = canvas.getBoundingClientRect();
        const scaleX = canvas.width / canvasRect.width;
        const scaleY = canvas.height / canvasRect.height;
        
        return {{
            x: (rect.left - canvasRect.left + rect.width / 2) * scaleX,
            y: (rect.top - canvasRect.top + rect.height / 2) * scaleY
        }};
    }}

    function drawArrow(fromX, fromY, toX, toY) {{
        const headLength = 10;
        const dx = toX - fromX;
        const dy = toY - fromY;
        const angle = Math.atan2(dy, dx);

        ctx.beginPath();
        ctx.moveTo(fromX, fromY);
        ctx.lineTo(toX, toY);
        ctx.stroke();

        ctx.beginPath();
        ctx.moveTo(toX, toY);
        ctx.lineTo(toX - headLength * Math.cos(angle - Math.PI / 6), toY - headLength * Math.sin(angle - Math.PI / 6));
        ctx.lineTo(toX - headLength * Math.cos(angle + Math.PI / 6), toY - headLength * Math.sin(angle + Math.PI / 6));
        ctx.closePath();
        ctx.fill();
    }}

    function deleteCard(cardId) {{
        updateStreamlit({{ type: 'delete', cardId: cardId }});
    }}

    function updateStreamlit(action) {{
        window.parent.postMessage({{
            type: "streamlit:setComponentValue",
            value: JSON.stringify({{ action: action }}),
        }}, "*");
    }}

    // Initialize
    resizeCanvas();
    initializeCards(cardData);
    drawConnections();
    updateMiniMap();

    // Panning functionality
    canvasContainer.addEventListener('mousedown', (e) => {{
        if (e.target === canvasContainer || e.target === canvas) {{
            isPanning = true;
            startX = e.clientX;
            startY = e.clientY;
        }}
    }});

    document.addEventListener('mousemove', onMouseMove);
    document.addEventListener('mouseup', stopDragging);

    // Mini-map interaction
    miniMapCanvas.addEventListener('mousedown', (e) => {{
        const rect = miniMapCanvas.getBoundingClientRect();
        const x = e.clientX - rect.left;
        const y = e.clientY - rect.top;
        const scaleX = containerWidth / miniMapCanvas.width;
        const scaleY = containerHeight / miniMapCanvas.height;
        const centerX = x * scaleX;
        const centerY = y * scaleY;
        
        offsetX = -centerX + outerContainer.offsetWidth / (2 * scale);
        offsetY = -centerY + outerContainer.offsetHeight / (2 * scale);
        
        cardContainer.style.left = offsetX + 'px';
        cardContainer.style.top = offsetY + 'px';
        
        updateMiniMap();
    }});

    // Mouse wheel zoom
    outerContainer.addEventListener('wheel', (e) => {{
        e.preventDefault();
        const delta = e.deltaY;
        if (delta > 0) {{
            zoomOut();
        }} else {{
            zoomIn();
        }}
    }});

    // Touch events for mobile devices
    let touchStartX, touchStartY;
    outerContainer.addEventListener('touchstart', (e) => {{
        if (e.touches.length === 1) {{
            isPanning = true;
            touchStartX = e.touches[0].clientX;
            touchStartY = e.touches[0].clientY;
        }}
    }});

    outerContainer.addEventListener('touchmove', (e) => {{
        if (e.touches.length === 1 && isPanning) {{
            e.preventDefault();
            const dx = e.touches[0].clientX - touchStartX;
            const dy = e.touches[0].clientY - touchStartY;
            offsetX += dx;
            offsetY += dy;
            cardContainer.style.left = offsetX + 'px';
            cardContainer.style.top = offsetY + 'px';
            touchStartX = e.touches[0].clientX;
            touchStartY = e.touches[0].clientY;
            updateMiniMap();
        }}
    }});

    outerContainer.addEventListener('touchend', () => {{
        isPanning = false;
    }});

    // Pinch-to-zoom for mobile devices
    let initialDistance = 0;
    outerContainer.addEventListener('touchstart', (e) => {{
        if (e.touches.length === 2) {{
            initialDistance = Math.hypot(
                e.touches[0].clientX - e.touches[1].clientX,
                e.touches[0].clientY - e.touches[1].clientY
            );
        }}
    }});

    outerContainer.addEventListener('touchmove', (e) => {{
        if (e.touches.length === 2) {{
            e.preventDefault();
            const currentDistance = Math.hypot(
                e.touches[0].clientX - e.touches[1].clientX,
                e.touches[0].clientY - e.touches[1].clientY
            );
            if (initialDistance > 0) {{
                if (currentDistance > initialDistance) {{
                    zoomIn();
                }} else {{
                    zoomOut();
                }}
            }}
            initialDistance = currentDistance;
        }}
    }});

    window.addEventListener('resize', drawConnections);
    cardContainer.addEventListener('scroll', drawConnections);

    // Initialize the board
    resizeCanvas();
    cardData.forEach(createCard);
    drawConnections();
    updateMiniMap();
    </script>
    """

    st.components.v1.html(html_content, height=820)

    if st.session_state.card_updates:
        try:
            data = json.loads(st.session_state.card_updates)
            action = data.get('action', {})
            
            if action.get('type') == 'delete':
                delete_card(action['cardId'])
            elif action.get('type') == 'update':
                update_card_position(action['cardId'], action['x'], action['y'])
            
            st.session_state.card_updates = None 
        except json.JSONDecodeError:
            st.error("Error processing card updates. Please try again.")
        except KeyError as e:
            st.error(f"Error accessing data: {str(e)}. Please check the data structure.")

def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def chatbot():
    st.subheader("Investigation Assistant")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    chat_container = st.container()
    input_container = st.container()

    with chat_container:
        for message in st.session_state.messages:
            st.text(f"{message['role'].capitalize()}: {message['content']}")

    with input_container:
        with st.form(key="chat_form"):
            prompt = st.text_input("What would you like to know about the case?", key="chat_input")
            col1, col2 = st.columns([1,1])
            with col1:
                submit_button = st.form_submit_button("Send")
            with col2:
                clear_button = st.form_submit_button("Clear Chat")

        if submit_button and prompt:
            st.session_state.messages.append({"role": "user", "content": prompt})

            context = "Case Information:\n"
            if "cards" in st.session_state and st.session_state.cards:
                for card in st.session_state.cards:
                    context += f"{card['type']} Card:\n"
                    for key, value in card.items():
                        if key not in ['id', 'x', 'y', 'image']:
                            context += f"  {key}: {value}\n"
                    context += "\n"
            else:
                context += "No cards have been added to the investigation board yet.\n"

            if "connections" in st.session_state and st.session_state.connections:
                context += "Connections:\n"
                for connection in st.session_state.connections:
                    context += f"  {connection[0]} is connected to {connection[1]}\n"

            ai_messages = [
                {"role": "system", "content": "You are a helpful assistant for a detective. Use the detailed case information to answer questions. If there's no relevant information in the case details, say so."},
                {"role": "user", "content": f"{context}\n\nQuestion: {prompt}"}
            ]

            try:
                with st.spinner("Thinking..."):
                    response = client.chat.completions.create(
                        model="tiiuae/falcon-180b-chat",
                        messages=ai_messages
                    )
                    full_response = response.choices[0].message.content
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
                full_response = "I'm sorry, but I encountered an error. Please try again later."

            st.session_state.messages.append({"role": "assistant", "content": full_response})
            
            st.experimental_rerun()

        if clear_button:
            st.session_state.messages = []
            st.experimental_rerun()

    st.write("Debug - Cards in session state:", st.session_state.get("cards", []))

def get_coordinates(location):
    geolocator = Nominatim(user_agent="detective_app")
    try:
        location_data = geolocator.geocode(location)
        if location_data:
            return location_data.latitude, location_data.longitude
        else:
            return None
    except (GeocoderTimedOut, GeocoderUnavailable):
        return None

def create_map_with_pins(cards, filter_option=None, filter_value=None, sort_by=None):
    locations = []
    for card in cards:
        if 'location' in card:
            coords = get_coordinates(card['location'])
            if coords:
                locations.append((coords[0], coords[1], card['type'], card.get('title', '') or card.get('name', ''), card))

    if not locations:
        st.warning("No valid locations found in the cards.")
        return None, []

    if filter_option and filter_value:
        locations = [loc for loc in locations if filter_value.lower() in str(loc[4].get(filter_option, '')).lower()]

    if sort_by:
        locations.sort(key=lambda x: x[4].get(sort_by, ''))

    if not locations:
        st.warning("No locations match the current filter and sort settings.")
        return None, []

    center_lat = sum(loc[0] for loc in locations) / len(locations)
    center_lon = sum(loc[1] for loc in locations) / len(locations)

    m = folium.Map(location=[center_lat, center_lon], zoom_start=10, width='100%', height='600px')

    for lat, lon, card_type, title, _ in locations:
        icon_color = 'red' if card_type == 'Event' else 'green' if card_type == 'Suspect' else 'blue'
        folium.Marker(
            [lat, lon],
            popup=f"{card_type}: {title}",
            icon=folium.Icon(color=icon_color, icon='info-sign')
        ).add_to(m)

    return m, locations

def connect_all_pins(m, locations):
    coordinates = [(loc[0], loc[1]) for loc in locations]
    folium.PolyLine(coordinates, color="purple", weight=2, opacity=0.8).add_to(m)
    return m

def handle_map_view():
    st.header("Map View")
    col1, col2 = st.columns([3, 1])
    
    with col2:
        filter_option = st.selectbox("Filter by:", ['None', 'type', 'title', 'name', 'date', 'location'])
        filter_value = st.text_input("Filter value:") if filter_option != 'None' else None
        sort_by = st.selectbox("Sort by:", ['None', 'type', 'title', 'name', 'date', 'location'])
        connect_pins = st.checkbox("Connect all pins")
    
    with col1:
        map, locations = create_map_with_pins(st.session_state.cards, filter_option, filter_value, sort_by)
        if map:
            if connect_pins:
                map = connect_all_pins(map, locations)
            folium_static(map, width=800, height=600)
        else:
            st.warning("No locations to display on the map.")
               
def get_lat_lon(location):
    geolocator = Nominatim(user_agent="ipdr_analysis")
    try:
        loc = geolocator.geocode(location)
        if loc:
            return loc.latitude, loc.longitude
        return None
    except (GeocoderTimedOut, GeocoderUnavailable):
        return None


def ipdr_analysis():
    st.subheader("IPDR Analysis")
    
    if 'ipdr_df' not in st.session_state:
        st.session_state.ipdr_df = None
    if 'ipdr_insights' not in st.session_state:
        st.session_state.ipdr_insights = None
    if 'ipdr_messages' not in st.session_state:
        st.session_state.ipdr_messages = []
    if 'ipdr_question' not in st.session_state:
        st.session_state.ipdr_question = ""
    if 'ipdr_context' not in st.session_state:
        st.session_state.ipdr_context = ""
    
    uploaded_file = st.file_uploader("Upload IPDR CSV file", type=['csv'])
    if uploaded_file is not None:
        st.session_state.ipdr_df = pd.read_csv(uploaded_file)
        df = st.session_state.ipdr_df
        
        ipdr_params = [
            "Calling Mobile Number", "Called Mobile Number", "Duration of Session",
            "Start Time", "End Time", "Amount of Data Transferred",
            "Internal IP Address", "External IP Address", "Port Number",
            "Cell Tower ID", "Cell Tower Location", "Azimuth Angle",
            "Protocol Used", "Service Used"
        ]
        parameter_mapping = {}
        for param in ipdr_params:
            best_match = difflib.get_close_matches(param, df.columns, n=1, cutoff=0.6)
            if best_match:
                parameter_mapping[param] = best_match[0]
            else:
                parameter_mapping[param] = None
        
        reverse_mapping = {v: k for k, v in parameter_mapping.items() if v}
        df = df.rename(columns=reverse_mapping)
        
        with st.expander("IPDR Data Preview", expanded=False):
            st.dataframe(df.head())

        
        st.markdown("""
        <style>
        .stGrid > div {
            background-color: #f0f2f6;
            border-radius: 10px;
            padding: 10px;
            margin: 10px 0;
        }
        </style>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)

        
        with col1:
            with st.container():
                options = {
                    "Frequently Called Numbers": ('Called Mobile Number', 'Times Called'),
                    "Frequently Calling Numbers": ('Calling Mobile Number', 'Times Called'),
                    "High Data Usage Sessions": ('Amount of Data Transferred', 'Data Transferred'),
                    "Long Duration Calls": ('Duration of Session', 'Duration'),
                    "Most Active Cell Towers": ('Cell Tower ID', 'Number of Connections')
                }
                
                selected_option = st.selectbox("Select List:", list(options.keys()))
                
                column, value_label = options[selected_option]
                
                if column in df.columns:
                    all_records = df[column].value_counts().reset_index()
                    all_records.columns = [column, value_label]
                    
                    # Pagination
                    records_per_page = 10
                    total_pages = len(all_records) // records_per_page + (1 if len(all_records) % records_per_page > 0 else 0)
                    page = st.number_input("Page", min_value=1, max_value=total_pages, value=1)
                    
                    start_idx = (page - 1) * records_per_page
                    end_idx = start_idx + records_per_page
                    
                    st.table(all_records.iloc[start_idx:end_idx])
                    st.write(f"Page {page} of {total_pages}")
                else:
                    st.warning(f"{column} not found in the dataset.")
        
        with col2:
            with st.container():
                st.markdown("### Value Mapping Graph")
                
                # Create a graph
                G = nx.Graph()
                
                # Add nodes and edges based on actual data
                for _, row in df.iterrows():
                    caller = row['Calling Mobile Number']
                    called = row['Called Mobile Number']
                    duration = row['Duration of Session']
                    G.add_edge(caller, called, weight=duration)
                
                # Limit the graph size for better performance
                if len(G.nodes()) > 100:
                    G = nx.k_core(G, 2)  # Keep only nodes with at least 2 connections
                
                # Create the plotly figure
                pos = nx.spring_layout(G)
                
                edge_x = []
                edge_y = []
                edge_text = []
                for edge in G.edges(data=True):
                    x0, y0 = pos[edge[0]]
                    x1, y1 = pos[edge[1]]
                    edge_x.extend([x0, x1, None])
                    edge_y.extend([y0, y1, None])
                    edge_text.append(f"Duration: {edge[2]['weight']}")

                edge_trace = go.Scatter(
                    x=edge_x, y=edge_y,
                    line=dict(width=0.5, color='#888'),
                    hoverinfo='text',
                    text=edge_text,
                    mode='lines')

                node_x = []
                node_y = []
                node_text = []
                node_adjacencies = []
                for node, adjacencies in G.adjacency():
                    x, y = pos[node]
                    node_x.append(x)
                    node_y.append(y)
                    node_adjacencies.append(len(adjacencies))
                    node_text.append(f'{node}<br>{len(adjacencies)}')

                node_trace = go.Scatter(
                    x=node_x, y=node_y,
                    mode='markers+text',
                    hoverinfo='none',
                    marker=dict(
                        showscale=True,
                        colorscale='YlGnBu',
                        reversescale=True,
                        color=node_adjacencies,
                        size=10,
                        colorbar=dict(
                            thickness=15,
                            title='Node Connections',
                            xanchor='left',
                            titleside='right'
                        ),
                        line_width=2),
                    text=node_text,
                    textposition="bottom center",
                    textfont=dict(size=8)
                )

                # Create the layout with zoom capabilities
                layout = go.Layout(
                    title='Value Mapping Graph',
                    titlefont_size=16,
                    showlegend=False,
                    hovermode='closest',
                    margin=dict(b=20,l=5,r=5,t=40),
                    annotations=[ dict(
                        text="Node color represents number of connections",
                        showarrow=False,
                        xref="paper", yref="paper",
                        x=0.005, y=-0.002 ) ],
                    xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                    yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                    dragmode='pan'
                )

                fig = go.Figure(data=[edge_trace, node_trace], layout=layout)

                # Add reset button
                fig.update_layout(
                    updatemenus=[
                        dict(
                            type="buttons",
                            direction="left",
                            buttons=[
                                dict(args=[{"xaxis.range": [-1.5, 1.5], "yaxis.range": [-1.5, 1.5]}],
                                     label="Reset View",
                                     method="relayout"
                                )
                            ],
                            pad={"r": 10, "t": 10},
                            showactive=False,
                            x=0.5,
                            xanchor="center",
                            y=-0.1,
                            yanchor="bottom"
                        ),
                    ]
                )

                # Add zoom and pan capabilities
                fig.update_layout(dragmode='pan')
                fig.update_xaxes(range=[-1.5, 1.5])
                fig.update_yaxes(range=[-1.5, 1.5])

                st.plotly_chart(fig, use_container_width=True)

                st.markdown("""
                **Graph Instructions:**
                - Pan: Click and drag
                - Zoom: Scroll or pinch
                - Reset: Use the 'Reset View' button below the graph
                - Node labels show: Phone number (top) and number of connections (bottom)
                - Edge hover: Shows call duration
                """)

        
        with col3:
            with st.container():
                st.markdown("### Geographic Heatmap")
                if "Cell Tower Location" in df.columns:
                    locations = df['Cell Tower Location'].unique()
                    geolocator = Nominatim(user_agent="ipdr_analysis")
                    lat_lons = []
                    for loc in locations:
                        try:
                            location = geolocator.geocode(loc)
                            if location:
                                lat_lons.append((location.latitude, location.longitude))
                        except (GeocoderTimedOut, GeocoderUnavailable):
                            pass
                    
                    if lat_lons:
                        m = folium.Map(location=lat_lons[0], zoom_start=6)
                        heatmap_data = [[lat, lon, 1] for lat, lon in lat_lons]
                        HeatMap(heatmap_data).add_to(m)
                        folium_static(m)
                    else:
                        st.warning("Could not geocode any locations for the heatmap.")
                else:
                    st.warning("Cell Tower Location column not found in the dataset.")
        
        with col1:
            with st.container():
                st.markdown("### Session Duration and Timing")
                if "Start Time" in df.columns:
                    df['Hour'] = pd.to_datetime(df['Start Time']).dt.hour
                    hourly_activity = df['Hour'].value_counts().sort_index()
                    fig = px.line(x=hourly_activity.index, y=hourly_activity.values, title="Hourly Activity")
                    fig.update_layout(height=300)
                    st.plotly_chart(fig)
                else:
                    st.warning("Start Time column not found in the dataset.")
                
                if "Duration of Session" in df.columns:
                    fig = px.histogram(df, x="Duration of Session", title="Call Duration Distribution", nbins=50)
                    fig.update_layout(height=300)
                    st.plotly_chart(fig)
                else:
                    st.warning("Duration of Session column not found in the dataset.")
        
        with col2:
            with st.container():
                st.markdown("### Data Usage Analysis")
                if "Amount of Data Transferred" in df.columns:
                    fig = px.histogram(df, x="Amount of Data Transferred", title="Data Usage Distribution", nbins=50)
                    fig.update_layout(height=300)
                    st.plotly_chart(fig)
                    
                    st.write(f"Total Data Transferred: {df['Amount of Data Transferred'].sum()} bytes")
                    st.write(f"Average Data per Session: {df['Amount of Data Transferred'].mean():.2f} bytes")
                else:
                    st.warning("Amount of Data Transferred column not found in the dataset.")
        
        with col3:
            with st.container():
                st.markdown("### IP Address Analysis")
                if "Internal IP Address" in df.columns and "External IP Address" in df.columns:
                    internal_ips = df['Internal IP Address'].nunique()
                    external_ips = df['External IP Address'].nunique()
                    st.write(f"Number of unique internal IPs: {internal_ips}")
                    st.write(f"Number of unique external IPs: {external_ips}")
                    
                    # Visualize IP address distribution
                    ip_counts = df['Internal IP Address'].value_counts().head(10)
                    fig = px.bar(x=ip_counts.index, y=ip_counts.values, title="Top 10 Internal IP Addresses")
                    fig.update_layout(height=300)
                    st.plotly_chart(fig)
                else:
                    st.warning("Internal IP Address or External IP Address columns not found in the dataset.")
        
        with col1:
            with st.container():
                st.markdown("### Cell Tower and Location Data")
                if "Azimuth Angle" in df.columns:
                    fig = px.histogram(df, x="Azimuth Angle", title="Azimuth Angle Distribution", nbins=36)
                    fig.update_layout(height=300)
                    st.plotly_chart(fig)
                else:
                    st.warning("Azimuth Angle column not found in the dataset.")
        
        with col2:
            with st.container():
                st.markdown("### Service and Protocol Usage")
                if "Service Used" in df.columns:
                    service_usage = df['Service Used'].value_counts()
                    fig = px.pie(values=service_usage.values, names=service_usage.index, title='Service Usage Distribution')
                    fig.update_layout(height=300)
                    st.plotly_chart(fig)
                else:
                    st.warning("Service Used column not found in the dataset.")
        
        with col3:
            with st.container():
                st.markdown("### Correlation Analysis")
                numeric_columns = df.select_dtypes(include=[np.number]).columns
                if len(numeric_columns) > 1:
                    correlation_matrix = df[numeric_columns].corr()
                    fig = px.imshow(correlation_matrix, text_auto=True, aspect="auto", title="Correlation Matrix")
                    fig.update_layout(height=400)
                    st.plotly_chart(fig)
                else:
                    st.warning("Not enough numeric columns for correlation analysis.")
        
        # Prepare context
        st.session_state.ipdr_context = f"""
        IPDR Data Summary:
        Total Records: {len(df)}
        Date Range: {df['Start Time'].min() if 'Start Time' in df.columns else 'N/A'} to {df['Start Time'].max() if 'Start Time' in df.columns else 'N/A'}
        Total Data Transferred: {df['Amount of Data Transferred'].sum() if 'Amount of Data Transferred' in df.columns else 'N/A'} bytes
        Unique Callers: {df['Calling Mobile Number'].nunique() if 'Calling Mobile Number' in df.columns else 'N/A'}
        Unique Called Numbers: {df['Called Mobile Number'].nunique() if 'Called Mobile Number' in df.columns else 'N/A'}
        Most Common Service: {df['Service Used'].mode().values[0] if 'Service Used' in df.columns else 'N/A'}
        Most Common Protocol: {df['Protocol Used'].mode().values[0] if 'Protocol Used' in df.columns else 'N/A'}
        Peak Activity Hour: {df.groupby(pd.to_datetime(df['Start Time']).dt.hour).size().idxmax() if 'Start Time' in df.columns else 'N/A'}
        Most Active Day: {df.groupby(pd.to_datetime(df['Start Time']).dt.day_name()).size().idxmax() if 'Start Time' in df.columns else 'N/A'}
        """
        
        # AI-Generated Insights
        st.markdown("### AI-Generated Insights")
        with st.container():
            if st.session_state.ipdr_insights is None or st.button("Regenerate Insights"):
                messages = [
                    {"role": "system", "content": "You are an expert in analyzing IPDR data. Provide insights based on the given data summary."},
                    {"role": "user", "content": f"{st.session_state.ipdr_context}\n\nProvide key insights and potential areas of interest based on this IPDR data summary."}
                ]
                
                with st.spinner("Generating AI insights..."):
                    try:
                        response = client.chat.completions.create(
                            model="tiiuae/falcon-180b-chat",
                            messages=messages
                        )
                        st.session_state.ipdr_insights = response.choices[0].message.content
                    except Exception as e:
                        st.error(f"An error occurred during insight generation: {str(e)}")
                        st.session_state.ipdr_insights = "Error generating insights."
            
            st.write("AI-Generated Insights:", st.session_state.ipdr_insights)
        
        # Generate PDF Report button
        if st.button("Generate PDF Report"):
            with st.spinner("Generating PDF report..."):
                pdf_buffer = generate_pdf_report(df, st.session_state.ipdr_insights)
                st.download_button(
                    label="Download PDF Report",
                    data=pdf_buffer,
                    file_name="ipdr_analysis_report.pdf",
                    mime="application/pdf"
                )
    else:
        st.warning("Please upload an IPDR CSV file to begin analysis.")
                    
def generate_pdf_report(df, insights):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
    from io import BytesIO
    import matplotlib.pyplot as plt
    import seaborn as sns

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph("IPDR Analysis Report", styles['Title']))
    story.append(Spacer(1, 12))

    # Summary Statistics
    story.append(Paragraph("Summary Statistics", styles['Heading2']))
    story.append(Paragraph(f"Total Records: {len(df)}", styles['Normal']))
    story.append(Paragraph(f"Date Range: {df['Start Time'].min()} to {df['Start Time'].max()}", styles['Normal']))
    story.append(Paragraph(f"Total Data Transferred: {df['Amount of Data Transferred'].sum()} bytes", styles['Normal']))
    story.append(Spacer(1, 12))

    # AI-Generated Insights
    story.append(Paragraph("AI-Generated Insights", styles['Heading2']))
    story.append(Paragraph(insights, styles['Normal']))
    story.append(Spacer(1, 12))

    # Visualizations
    story.append(Paragraph("Data Visualizations", styles['Heading2']))
    
    # Hourly Activity
    plt.figure(figsize=(8, 4))
    df['Hour'] = pd.to_datetime(df['Start Time']).dt.hour
    hourly_activity = df['Hour'].value_counts().sort_index()
    sns.lineplot(x=hourly_activity.index, y=hourly_activity.values)
    plt.title("Hourly Activity")
    plt.xlabel("Hour of Day")
    plt.ylabel("Number of Sessions")
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png')
    img_buffer.seek(0)
    story.append(Image(img_buffer, width=400, height=200))
    story.append(Spacer(1, 12))
    
    # Call Duration Distribution
    plt.figure(figsize=(8, 4))
    sns.histplot(df['Duration of Session'], bins=50)
    plt.title("Call Duration Distribution")
    plt.xlabel("Duration (seconds)")
    plt.ylabel("Frequency")
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png')
    img_buffer.seek(0)
    story.append(Image(img_buffer, width=400, height=200))
    
    # Build the PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_key_statistics(df):
    return {
        "total_records": len(df),
        "date_range": {
            "start": df['Start Time'].min(),
            "end": df['Start Time'].max()
        },
        "total_data_transferred": df['Amount of Data Transferred'].sum(),
        "unique_callers": df['Calling Mobile Number'].nunique(),
        "unique_called_numbers": df['Called Mobile Number'].nunique(),
        "most_common_service": df['Service Used'].mode().values[0],
        "most_common_protocol": df['Protocol Used'].mode().values[0],
        "peak_activity_hour": df.groupby(pd.to_datetime(df['Start Time']).dt.hour).size().idxmax(),
        "most_active_day": df.groupby(pd.to_datetime(df['Start Time']).dt.day_name()).size().idxmax()
    }

exa = Exa(EXA_API_KEY)

def generate_search_query(question):
    SYSTEM_MESSAGE = "You are a helpful assistant that generates search queries based on user questions. Only generate one search query."
    
    try:
        completion = client.chat.completions.create(
            model="tiiuae/falcon-180b-chat",
            messages=[
                {"role": "system", "content": SYSTEM_MESSAGE},
                {"role": "user", "content": question},
            ],
        )
        search_query = completion.choices[0].message.content
        return search_query
    except Exception as e:
        st.error(f"Error generating search query: {str(e)}")
        return None

def search_internet(question):
    search_query = generate_search_query(question)
    if not search_query:
        return None, None

    one_week_ago = (datetime.now() - timedelta(days=7))
    date_cutoff = one_week_ago.strftime("%Y-%m-%d")

    try:
        search_response = exa.search_and_contents(
            search_query, use_autoprompt=True, start_published_date=date_cutoff
        )
        
        urls = [result.url for result in search_response.results]
        contents = [result.text for result in search_response.results]
        
        return urls, contents
    except Exception as e:
        st.error(f"Error searching the internet: {str(e)}")
        return None, None

def summarize_search_results(question, urls, contents):
    if not urls or not contents:
        return "Sorry, I couldn't find any relevant information."

    context = "\n\n".join([f"URL: {url}\nContent: {content[:500]}..." for url, content in zip(urls[:3], contents[:3])])
    
    prompt = f"""Based on the following search results, please provide a concise summary answering the question: "{question}"
    
    Search Results:
    {context}
    
    Summary:"""

    try:
        response = client.chat.completions.create(
            model="tiiuae/falcon-180b-chat",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that summarizes search results to answer user questions. Provide a concise summary and always cite the sources using [1], [2], etc."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error summarizing search results: {str(e)}")
        return "Sorry, I encountered an error while summarizing the search results."

def internet_search_interface():
    st.subheader("Internet Search")
    
    query = st.text_input("Enter your question:")
    if st.button("Search"):
        with st.spinner("Searching the internet..."):
            urls, contents = search_internet(query)
            if urls and contents:
                summary = summarize_search_results(query, urls, contents)
                st.write("Summary:")
                st.write(summary)
                st.write("Sources:")
                for i, url in enumerate(urls[:3], 1):
                    st.write(f"[{i}] {url}")
            else:
                st.warning("No results found or an error occurred during the search.")


def main():
    st.set_page_config(layout="wide")
    st.title("InvestiBoard")

    initialize_session_state()
    
    if st.session_state.authentication_status != True:
        tab1, tab2 = st.tabs(["Sign In", "Sign Up"])
        
        with tab1:
            sign_in()
        
        with tab2:
            sign_up()
    
    else:
        st.sidebar.write(f"Welcome, {st.session_state.user['email']}!")
        if st.sidebar.button("Sign Out"):
            sign_out()
            st.experimental_rerun()
        
        tab1, tab2 = st.tabs(["Investigation Board", "Analysis Tools"])
        
        with tab1:
            col1, col2 = st.columns([3, 1])
            
            with col1:
                st.header("Investigation Board")
                display_draggable_cards()
            
            with col2:
                with st.expander("Add New Card", expanded=False):
                    card_type = st.selectbox("Select card type", ["Event", "Suspect", "Evidence"])
                    create_card(card_type)
                
                with st.expander("Connect Cards", expanded=False):
                    connect_cards()
                
                chatbot()
                
                if st.button("Clear All Cards"):
                    clear_all_cards()
                    st.success("All cards have been cleared.")
                    st.experimental_rerun()

        
        with tab2:
            tool_tabs = st.tabs(["Map View", "Tabular View", "IPFS", "RAG", "Internet Search"])
            
            with tool_tabs[0]:  
                st.header("Map View")
                col1, col2 = st.columns([3, 1])
                
                with col2:
                    filter_option = st.selectbox("Filter by:", ['None', 'type', 'title', 'name', 'date', 'location'])
                    filter_value = st.text_input("Filter value:") if filter_option != 'None' else None
                    sort_by = st.selectbox("Sort by:", ['None', 'type', 'title', 'name', 'date', 'location'])
                    connect_pins = st.checkbox("Connect all pins")
                
                with col1:
                    map, locations = create_map_with_pins(st.session_state.cards, filter_option, filter_value, sort_by)
                    if map:
                        if connect_pins:
                            map = connect_all_pins(map, locations)
                        folium_static(map, width=800, height=600)
                    else:
                        st.warning("No locations to display on the map.")
            
            with tool_tabs[1]: 
                st.header("Tabular View")
                if st.session_state.cards:
                    df = create_tabular_view(st.session_state.cards)
                    
                    col1, col2, col3 = st.columns([1, 1, 1])
                    with col1:
                        sort_options = {
                            'Time (ascending)': ('time', True),
                            'Time (descending)': ('time', False),
                            'Date (ascending)': ('date', True),
                            'Date (descending)': ('date', False),
                            'Location': ('location', True),
                            'Suspect Name': ('name', True),
                            'Evidence Type': ('evidence_type', True)
                        }
                        sort_by = st.selectbox("Sort by:", list(sort_options.keys()))
                    
                    with col2:
                        filter_options = ['None', 'Location', 'Suspect Name', 'Evidence Related']
                        filter_option = st.selectbox("Filter by:", filter_options)
                    
                    with col3:
                        filter_value = None
                        if filter_option != 'None':
                            filter_value = st.text_input(f"Enter {filter_option.lower()} to filter:")
                    
                    filter_map = {
                        'Location': 'location',
                        'Suspect Name': 'suspect_name',
                        'Evidence Related': 'evidence_related'
                    }
                    
                    df_display = display_tabular_view(
                        df, 
                        sort_by=sort_options[sort_by][0],
                        ascending=sort_options[sort_by][1],
                        filter_option=filter_map.get(filter_option),
                        filter_value=filter_value
                    )
                    
                    st.dataframe(df_display, use_container_width=True)
                    
                    if st.button("Export to CSV"):
                        csv = df_display.to_csv(index=False)
                        b64 = base64.b64encode(csv.encode()).decode()
                        href = f'<a href="data:file/csv;base64,{b64}" download="investigation_data.csv">Download CSV File</a>'
                        st.markdown(href, unsafe_allow_html=True)
                else:
                    st.info("No cards available for tabular view.")
            
            with tool_tabs[2]:  
                st.header("IPFS File Management")
                
                if not st.session_state.get('wallet_connected', False):
                    connect_wallet()
                else:
                    st.success(f"Wallet connected: {st.session_state.wallet_address}")
                    if st.button("Disconnect Wallet"):
                        disconnect_wallet()
                        st.experimental_rerun()
                
                if st.session_state.get('wallet_connected', False):
                    ipfs_tabs = st.tabs(["Upload File", "View Files"])
                    
                    with ipfs_tabs[0]:
                        st.subheader("Upload File to IPFS")
                        uploaded_file = st.file_uploader("Choose a file to upload", type=['pdf', 'docx', 'txt', 'jpg', 'png'])
                        if uploaded_file is not None:
                            if st.button("Upload to IPFS"):
                                with st.spinner("Uploading file to IPFS..."):
                                    ipfs_hash = upload_to_ipfs(uploaded_file)
                                    if ipfs_hash:
                                        st.success(f"File uploaded to IPFS. Hash: {ipfs_hash}")
                                        user_id = st.session_state.user['localId']
                                        db.child("users").child(user_id).child("ipfs_files").push({"name": uploaded_file.name, "hash": ipfs_hash})
                    
                    with ipfs_tabs[1]:
                        st.subheader("View IPFS Files")
                        user_id = st.session_state.user['localId']
                        user_files = db.child("users").child(user_id).child("ipfs_files").get().val()
                        if user_files:
                            for file_key, file_data in user_files.items():
                                st.write(f"File: {file_data['name']}")
                                st.write(f"IPFS Hash: {file_data['hash']}")
                                st.markdown(f"[View on IPFS](https://gateway.pinata.cloud/ipfs/{file_data['hash']})")
                        else:
                            st.info("No files uploaded yet.")
                else:
                    st.warning("Please connect your wallet to access IPFS features.")
            
            with tool_tabs[3]:  
                st.header("RAG & IPDR Analysis")
                rag_interface()
            
            with tool_tabs[4]:
                internet_search_interface()


if __name__ == "__main__":
    main()

